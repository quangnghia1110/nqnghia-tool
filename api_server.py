"""
API Server for NQNghia Tools.
Provides REST endpoints for all tools.

Run: python api_server.py
Server starts at http://localhost:8000
"""

import os
import io
import json
import uuid
import shutil
import subprocess
import urllib.parse
import webbrowser
import threading
import mimetypes
import tempfile
import re as _re_mod

# Fix numba cache issue on Windows Store Python
os.environ.setdefault("NUMBA_CACHE_DIR", os.path.join(tempfile.gettempdir(), "numba_cache"))
os.makedirs(os.environ["NUMBA_CACHE_DIR"], exist_ok=True)
from http.server import HTTPServer, SimpleHTTPRequestHandler
from video_downloader import download_video, detect_platform, get_progress, _progress_store

PORT = int(os.environ.get("PORT", 8000))
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOWNLOAD_DIR = os.path.join(BASE_DIR, "downloads")
TEMP_DIR = os.path.join(BASE_DIR, "temp")

# Shared stores
_task_results = {}   # task_id -> result dict
_file_store = {}     # file_id -> file_path
_convert_progress = {}  # task_id -> progress dict


def _parse_multipart(body: bytes, content_type: str):
    """Parse multipart/form-data without the deprecated cgi module.
    Returns (fields_dict, files_list) where files_list items are (fieldname, filename, data).
    """
    boundary = content_type.split("boundary=")[-1].strip()
    parts = body.split(b"--" + boundary.encode())
    fields = {}
    files = []
    for part in parts:
        part = part.strip()
        if not part or part == b"--":
            continue
        if b"\r\n\r\n" in part:
            header_section, data = part.split(b"\r\n\r\n", 1)
        elif b"\n\n" in part:
            header_section, data = part.split(b"\n\n", 1)
        else:
            continue
        # Remove trailing \r\n-- from data
        if data.endswith(b"\r\n"):
            data = data[:-2]
        headers_text = header_section.decode("utf-8", errors="replace")
        # Parse Content-Disposition
        name = None
        filename = None
        for line in headers_text.split("\n"):
            line = line.strip()
            if line.lower().startswith("content-disposition:"):
                name_match = _re_mod.search(r'name="([^"]*)"', line)
                fname_match = _re_mod.search(r'filename="([^"]*)"', line)
                if name_match:
                    name = name_match.group(1)
                if fname_match:
                    filename = fname_match.group(1)
        if name is None:
            continue
        if filename is not None:
            files.append((name, filename, data))
        else:
            fields[name] = data.decode("utf-8", errors="replace")
    return fields, files


class ToolsHandler(SimpleHTTPRequestHandler):
    """Handle API requests for all tools."""

    def do_GET(self):
        parsed = urllib.parse.urlparse(self.path)
        params = urllib.parse.parse_qs(parsed.query)

        if parsed.path == "/api/detect":
            self._handle_detect(params)
        elif parsed.path == "/api/progress":
            self._handle_progress(params)
        elif parsed.path == "/api/file":
            self._handle_file_serve(params)
        elif parsed.path == "/api/convert/progress":
            self._handle_convert_progress(params)
        elif parsed.path == "/api/convert/file":
            self._handle_file_serve(params)
        elif parsed.path == "/api/image-convert/progress":
            self._handle_convert_progress(params)
        elif parsed.path == "/api/image-convert/file":
            self._handle_file_serve(params)
        elif parsed.path == "/api/file-convert/progress":
            self._handle_convert_progress(params)
        elif parsed.path == "/api/file-convert/file":
            self._handle_file_serve(params)
        elif parsed.path == "/api/remove-bg/progress":
            self._handle_convert_progress(params)
        elif parsed.path == "/api/remove-bg/file":
            self._handle_file_serve(params)
        elif parsed.path == "/api/pdf/progress":
            self._handle_convert_progress(params)
        elif parsed.path == "/api/pdf/file":
            self._handle_file_serve(params)
        elif parsed.path == "/api/status":
            self._send_json({"status": "ok", "message": "NQNghia Tools API is running"})
        else:
            super().do_GET()

    def do_POST(self):
        parsed = urllib.parse.urlparse(self.path)

        if parsed.path == "/api/download/start":
            data = self._read_json_body()
            if data is None:
                return
            self._handle_download_start(data)
        elif parsed.path == "/api/detect":
            data = self._read_json_body()
            if data is None:
                return
            url = data.get("url", "")
            platform = detect_platform(url)
            self._send_json({"platform": platform, "url": url})
        elif parsed.path == "/api/convert":
            self._handle_convert()
        elif parsed.path == "/api/image-convert":
            self._handle_image_convert()
        elif parsed.path == "/api/file-convert":
            self._handle_file_convert()
        elif parsed.path == "/api/remove-bg":
            self._handle_remove_bg()
        elif parsed.path == "/api/audio/cut":
            self._handle_audio_cut()
        elif parsed.path == "/api/pdf/merge":
            self._handle_pdf_merge()
        elif parsed.path == "/api/pdf/split":
            self._handle_pdf_split()
        elif parsed.path == "/api/speed-test/upload":
            self._handle_speed_upload()
        else:
            self._send_json({"error": "Not found"}, status=404)

    # ===== Helpers =====

    def _read_json_body(self):
        content_length = int(self.headers.get("Content-Length", 0))
        body = self.rfile.read(content_length).decode("utf-8")
        try:
            return json.loads(body)
        except json.JSONDecodeError:
            self._send_json({"success": False, "error": "Invalid JSON"}, status=400)
            return None

    def _read_multipart(self):
        """Parse multipart/form-data, return (fields_dict, file_data, file_name)."""
        content_type = self.headers.get("Content-Type", "")
        if "multipart/form-data" not in content_type:
            self._send_json({"error": "Expected multipart/form-data"}, status=400)
            return None, None, None

        content_length = int(self.headers.get("Content-Length", 0))
        body = self.rfile.read(content_length)

        fields, files = _parse_multipart(body, content_type)

        file_data = None
        file_name = None
        if files:
            _, file_name, file_data = files[0]

        return fields, file_data, file_name

    # ===== Video Downloader =====

    def _handle_detect(self, params):
        url = params.get("url", [""])[0]
        if not url:
            self._send_json({"error": "Missing 'url' parameter"}, status=400)
            return
        platform = detect_platform(url)
        self._send_json({"platform": platform, "url": url})

    def _handle_progress(self, params):
        task_id = params.get("task_id", [""])[0]
        if not task_id:
            self._send_json({"error": "Missing 'task_id'"}, status=400)
            return
        progress = get_progress(task_id)
        if task_id in _task_results:
            result = _task_results.pop(task_id)
            _progress_store.pop(task_id, None)

            if result.get("success") and result.get("file_path"):
                file_id = str(uuid.uuid4())[:8]
                _file_store[file_id] = result["file_path"]
                file_name = os.path.basename(result["file_path"])
                result["file_id"] = file_id
                result["file_name"] = file_name

            progress = {**progress, "status": "done", "percent": 100, "result": result}
        self._send_json(progress)

    def _handle_download_start(self, data):
        url = data.get("url", "")
        if not url:
            self._send_json({"success": False, "error": "Missing 'url'"}, status=400)
            return
        quality = data.get("quality", "best")
        fmt = data.get("format", "mp4")
        task_id = str(uuid.uuid4())[:8]

        def _run():
            result = download_video(
                url, output_dir=DOWNLOAD_DIR, quality=quality,
                task_id=task_id, output_format=fmt,
            )
            _task_results[task_id] = result

        threading.Thread(target=_run, daemon=True).start()
        self._send_json({"task_id": task_id, "message": "Download started"})

    # ===== Media Converter (video) =====

    def _handle_convert(self):
        fields, file_data, file_name = self._read_multipart()
        if file_data is None:
            self._send_json({"error": "Không có file"}, status=400)
            return

        output_format = fields.get("output_format", "mp4")
        task_id = str(uuid.uuid4())[:8]
        _convert_progress[task_id] = {"status": "processing", "percent": 0}

        # Save uploaded file
        os.makedirs(TEMP_DIR, exist_ok=True)
        input_path = os.path.join(TEMP_DIR, f"{task_id}_{file_name}")
        with open(input_path, "wb") as f:
            f.write(file_data)

        out_name = os.path.splitext(file_name)[0] + f".{output_format}"
        output_path = os.path.join(TEMP_DIR, f"{task_id}_{out_name}")

        def _run():
            ffmpeg = shutil.which("ffmpeg")
            if not ffmpeg:
                _convert_progress[task_id] = {"status": "error", "error": "ffmpeg không được cài đặt"}
                return
            try:
                _convert_progress[task_id] = {"status": "processing", "percent": 50}
                subprocess.run(
                    [ffmpeg, "-i", input_path, "-y", output_path],
                    capture_output=True, timeout=600,
                )
                if os.path.exists(output_path):
                    file_id = str(uuid.uuid4())[:8]
                    _file_store[file_id] = output_path
                    _convert_progress[task_id] = {
                        "status": "done", "percent": 100,
                        "file_id": file_id, "file_name": out_name,
                        "file_size": os.path.getsize(output_path),
                    }
                else:
                    _convert_progress[task_id] = {"status": "error", "error": "Chuyển đổi thất bại"}
            except subprocess.TimeoutExpired:
                _convert_progress[task_id] = {"status": "error", "error": "Quá thời gian xử lý"}
            except Exception as e:
                _convert_progress[task_id] = {"status": "error", "error": str(e)}
            finally:
                _cleanup_file(input_path)

        threading.Thread(target=_run, daemon=True).start()
        self._send_json({"task_id": task_id})

    def _handle_convert_progress(self, params):
        task_id = params.get("task_id", [""])[0]
        if not task_id:
            self._send_json({"error": "Missing task_id"}, status=400)
            return
        progress = _convert_progress.get(task_id, {"status": "unknown"})
        self._send_json(progress)

    # ===== Remove Background =====

    def _handle_remove_bg(self):
        fields, file_data, file_name = self._read_multipart()
        if file_data is None:
            self._send_json({"error": "Không có file"}, status=400)
            return

        task_id = str(uuid.uuid4())[:8]
        _convert_progress[task_id] = {"status": "processing", "percent": 0}

        os.makedirs(TEMP_DIR, exist_ok=True)
        input_path = os.path.join(TEMP_DIR, f"{task_id}_{file_name}")
        with open(input_path, "wb") as f:
            f.write(file_data)

        out_name = os.path.splitext(file_name)[0] + "_nobg.png"
        output_path = os.path.join(TEMP_DIR, f"{task_id}_{out_name}")

        def _run():
            try:
                from rembg import remove
                from PIL import Image

                _convert_progress[task_id] = {"status": "processing", "percent": 30}

                with open(input_path, "rb") as inp:
                    input_data = inp.read()

                _convert_progress[task_id] = {"status": "processing", "percent": 60}
                output_data = remove(input_data)

                with open(output_path, "wb") as out:
                    out.write(output_data)

                if os.path.exists(output_path):
                    file_id = str(uuid.uuid4())[:8]
                    _file_store[file_id] = output_path
                    _convert_progress[task_id] = {
                        "status": "done", "percent": 100,
                        "file_id": file_id, "file_name": out_name,
                        "file_size": os.path.getsize(output_path),
                    }
                else:
                    _convert_progress[task_id] = {"status": "error", "error": "Xoá nền thất bại"}

            except ImportError:
                _convert_progress[task_id] = {
                    "status": "error",
                    "error": "Cần cài thư viện: pip install rembg[gpu] hoặc pip install rembg"
                }
            except Exception as e:
                _convert_progress[task_id] = {"status": "error", "error": str(e)}
            finally:
                _cleanup_file(input_path)

        threading.Thread(target=_run, daemon=True).start()
        self._send_json({"task_id": task_id})

    # ===== Audio Cutter =====

    def _handle_audio_cut(self):
        fields, file_data, file_name = self._read_multipart()
        if file_data is None:
            self._send_json({"error": "Không có file"}, status=400)
            return

        start = fields.get("start", "0")
        end = fields.get("end", "")
        out_format = fields.get("format", "mp3")

        os.makedirs(TEMP_DIR, exist_ok=True)
        task_id = str(uuid.uuid4())[:8]
        input_path = os.path.join(TEMP_DIR, f"{task_id}_{file_name}")
        with open(input_path, "wb") as f:
            f.write(file_data)

        out_name = os.path.splitext(file_name)[0] + f"_cut.{out_format}"
        output_path = os.path.join(TEMP_DIR, f"{task_id}_{out_name}")

        ffmpeg = shutil.which("ffmpeg")
        if not ffmpeg:
            _cleanup_file(input_path)
            self._send_json({"success": False, "error": "ffmpeg không được cài đặt"})
            return

        try:
            cmd = [ffmpeg, "-i", input_path, "-ss", start]
            if end:
                cmd += ["-to", end]
            if out_format == "mp3":
                cmd += ["-acodec", "libmp3lame", "-ab", "320k"]
            cmd += ["-y", output_path]

            subprocess.run(cmd, capture_output=True, timeout=120)

            if os.path.exists(output_path):
                file_id = str(uuid.uuid4())[:8]
                _file_store[file_id] = output_path
                self._send_json({
                    "success": True,
                    "file_id": file_id,
                    "file_name": out_name,
                    "file_size": os.path.getsize(output_path),
                })
            else:
                self._send_json({"success": False, "error": "Cắt audio thất bại"})
        except Exception as e:
            self._send_json({"success": False, "error": str(e)})
        finally:
            _cleanup_file(input_path)

    # ===== Image Converter (Pillow) =====

    def _handle_image_convert(self):
        fields, file_data, file_name = self._read_multipart()
        if file_data is None:
            self._send_json({"error": "Không có file"}, status=400)
            return

        output_format = fields.get("output_format", "png").lower()
        task_id = str(uuid.uuid4())[:8]
        _convert_progress[task_id] = {"status": "processing", "percent": 0}

        os.makedirs(TEMP_DIR, exist_ok=True)
        input_path = os.path.join(TEMP_DIR, f"{task_id}_{file_name}")
        with open(input_path, "wb") as f:
            f.write(file_data)

        def _run():
            try:
                from PIL import Image
                _convert_progress[task_id] = {"status": "processing", "percent": 30}

                # Map format names to Pillow format strings
                fmt_map = {
                    "jpg": "JPEG", "jpeg": "JPEG", "png": "PNG", "bmp": "BMP",
                    "gif": "GIF", "tiff": "TIFF", "tif": "TIFF", "webp": "WEBP",
                    "ico": "ICO", "eps": "EPS", "tga": "TGA",
                    "svg": "SVG", "wbmp": "WBMP",
                }
                pil_format = fmt_map.get(output_format, output_format.upper())

                out_name = os.path.splitext(file_name)[0] + f".{output_format}"
                output_path = os.path.join(TEMP_DIR, f"{task_id}_{out_name}")

                _convert_progress[task_id] = {"status": "processing", "percent": 50}

                img = Image.open(input_path)

                # Handle SVG output separately (rasterize not supported, use cairosvg or skip)
                if output_format == "svg":
                    _convert_progress[task_id] = {
                        "status": "error",
                        "error": "Chuyển đổi sang SVG cần file vector gốc, không hỗ trợ từ ảnh raster"
                    }
                    return

                # WBMP: convert to 1-bit then save as BMP (WBMP ~ 1-bit bitmap)
                if output_format == "wbmp":
                    img = img.convert("1")
                    pil_format = "BMP"
                    out_name = os.path.splitext(file_name)[0] + ".wbmp"
                    output_path = os.path.join(TEMP_DIR, f"{task_id}_{out_name}")

                # Handle RGBA -> RGB for formats that don't support alpha
                if pil_format in ("JPEG", "BMP", "EPS", "TGA") and img.mode in ("RGBA", "P"):
                    bg = Image.new("RGB", img.size, (255, 255, 255))
                    if img.mode == "P":
                        img = img.convert("RGBA")
                    bg.paste(img, mask=img.split()[3])
                    img = bg

                # ICO needs specific sizes
                if pil_format == "ICO":
                    sizes = [(256, 256)]
                    img.save(output_path, format="ICO", sizes=sizes)
                else:
                    save_kwargs = {}
                    if pil_format == "JPEG":
                        save_kwargs["quality"] = 95
                    elif pil_format == "WEBP":
                        save_kwargs["quality"] = 90
                    img.save(output_path, format=pil_format, **save_kwargs)

                if os.path.exists(output_path):
                    file_id = str(uuid.uuid4())[:8]
                    _file_store[file_id] = output_path
                    _convert_progress[task_id] = {
                        "status": "done", "percent": 100,
                        "file_id": file_id, "file_name": out_name,
                        "file_size": os.path.getsize(output_path),
                    }
                else:
                    _convert_progress[task_id] = {"status": "error", "error": "Chuyển đổi thất bại"}

            except ImportError:
                _convert_progress[task_id] = {
                    "status": "error",
                    "error": "Cần cài thư viện: pip install Pillow"
                }
            except Exception as e:
                _convert_progress[task_id] = {"status": "error", "error": str(e)}
            finally:
                _cleanup_file(input_path)

        threading.Thread(target=_run, daemon=True).start()
        self._send_json({"task_id": task_id})

    # ===== File/Document Converter =====

    def _handle_file_convert(self):
        fields, file_data, file_name = self._read_multipart()
        if file_data is None:
            self._send_json({"error": "Không có file"}, status=400)
            return

        output_format = fields.get("output_format", fields.get("mode", "pdf"))
        # Support old "mode" field for backwards compat
        if output_format == "pdf-to-word":
            output_format = "docx"
        elif output_format == "word-to-pdf":
            output_format = "pdf"

        task_id = str(uuid.uuid4())[:8]
        _convert_progress[task_id] = {"status": "processing", "percent": 0}

        os.makedirs(TEMP_DIR, exist_ok=True)
        input_path = os.path.join(TEMP_DIR, f"{task_id}_{file_name}")
        with open(input_path, "wb") as f:
            f.write(file_data)

        input_ext = os.path.splitext(file_name)[1].lower().lstrip(".")

        def _run():
            try:
                _convert_progress[task_id] = {"status": "processing", "percent": 20}

                out_name = os.path.splitext(file_name)[0] + f".{output_format}"
                output_path = os.path.join(TEMP_DIR, f"{task_id}_{out_name}")

                converted = False

                # Special case: PDF -> DOCX using pdf2docx
                if input_ext == "pdf" and output_format == "docx":
                    try:
                        from pdf2docx import Converter
                        _convert_progress[task_id] = {"status": "processing", "percent": 50}
                        cv = Converter(input_path)
                        cv.convert(output_path)
                        cv.close()
                        converted = True
                    except ImportError:
                        pass  # Fall through to LibreOffice

                # Special case: DOCX -> PDF using docx2pdf
                if not converted and input_ext in ("docx", "doc") and output_format == "pdf":
                    try:
                        from docx2pdf import convert
                        _convert_progress[task_id] = {"status": "processing", "percent": 50}
                        convert(input_path, output_path)
                        converted = True
                    except ImportError:
                        pass  # Fall through to LibreOffice

                # Special case: DOCX -> HTML using python-docx
                if not converted and input_ext in ("docx",) and output_format == "html":
                    try:
                        import docx
                        _convert_progress[task_id] = {"status": "processing", "percent": 50}
                        doc = docx.Document(input_path)
                        html_parts = [
                            "<!DOCTYPE html>",
                            "<html><head><meta charset='utf-8'></head><body>",
                        ]
                        for para in doc.paragraphs:
                            style = para.style.name if para.style else ""
                            if "Heading 1" in style:
                                html_parts.append(f"<h1>{para.text}</h1>")
                            elif "Heading 2" in style:
                                html_parts.append(f"<h2>{para.text}</h2>")
                            elif "Heading 3" in style:
                                html_parts.append(f"<h3>{para.text}</h3>")
                            elif para.text.strip():
                                # Preserve bold/italic
                                runs_html = ""
                                for run in para.runs:
                                    t = run.text
                                    if not t:
                                        continue
                                    if run.bold and run.italic:
                                        t = f"<b><i>{t}</i></b>"
                                    elif run.bold:
                                        t = f"<b>{t}</b>"
                                    elif run.italic:
                                        t = f"<i>{t}</i>"
                                    runs_html += t
                                html_parts.append(f"<p>{runs_html or para.text}</p>")
                        # Tables
                        for table in doc.tables:
                            html_parts.append("<table border='1' style='border-collapse:collapse'>")
                            for row in table.rows:
                                html_parts.append("<tr>")
                                for cell in row.cells:
                                    html_parts.append(f"<td style='padding:4px 8px'>{cell.text}</td>")
                                html_parts.append("</tr>")
                            html_parts.append("</table>")
                        html_parts.append("</body></html>")
                        with open(output_path, "w", encoding="utf-8") as hf:
                            hf.write("\n".join(html_parts))
                        converted = True
                    except ImportError:
                        pass

                # Special case: DOCX -> TXT
                if not converted and input_ext in ("docx",) and output_format == "txt":
                    try:
                        import docx
                        _convert_progress[task_id] = {"status": "processing", "percent": 50}
                        doc = docx.Document(input_path)
                        text = "\n".join(p.text for p in doc.paragraphs)
                        with open(output_path, "w", encoding="utf-8") as tf:
                            tf.write(text)
                        converted = True
                    except ImportError:
                        pass

                # General: use LibreOffice for all document conversions
                if not converted:
                    lo = shutil.which("soffice") or shutil.which("libreoffice")
                    if not lo:
                        _convert_progress[task_id] = {
                            "status": "error",
                            "error": "Cần cài LibreOffice để chuyển đổi tài liệu. Tải tại: https://www.libreoffice.org"
                        }
                        return

                    _convert_progress[task_id] = {"status": "processing", "percent": 50}

                    # LibreOffice format mapping
                    lo_format_map = {
                        "doc": "doc",
                        "docx": "docx",
                        "pdf": "pdf",
                        "html": "html",
                        "odt": "odt",
                        "ppt": "ppt",
                        "pptx": "pptx",
                        "rtf": "rtf",
                        "txt": "txt",
                        "xlsx": "xlsx",
                        "csv": "csv",
                    }
                    lo_fmt = lo_format_map.get(output_format, output_format)

                    result = subprocess.run(
                        [lo, "--headless", "--convert-to", lo_fmt,
                         "--outdir", TEMP_DIR, input_path],
                        capture_output=True, timeout=180,
                    )

                    # LibreOffice outputs with the original filename stem + new ext
                    lo_out = os.path.join(
                        TEMP_DIR,
                        os.path.splitext(f"{task_id}_{file_name}")[0] + f".{output_format}"
                    )
                    if os.path.exists(lo_out) and lo_out != output_path:
                        os.rename(lo_out, output_path)
                    converted = os.path.exists(output_path)

                if os.path.exists(output_path):
                    file_id = str(uuid.uuid4())[:8]
                    _file_store[file_id] = output_path
                    _convert_progress[task_id] = {
                        "status": "done", "percent": 100,
                        "file_id": file_id, "file_name": out_name,
                        "file_size": os.path.getsize(output_path),
                    }
                else:
                    _convert_progress[task_id] = {"status": "error", "error": "Chuyển đổi thất bại"}
            except subprocess.TimeoutExpired:
                _convert_progress[task_id] = {"status": "error", "error": "Quá thời gian xử lý"}
            except Exception as e:
                _convert_progress[task_id] = {"status": "error", "error": str(e)}
            finally:
                _cleanup_file(input_path)

        threading.Thread(target=_run, daemon=True).start()
        self._send_json({"task_id": task_id})

    # ===== PDF Tools =====

    def _handle_pdf_merge(self):
        """Merge multiple PDF files into one."""
        content_type = self.headers.get("Content-Type", "")
        if "multipart/form-data" not in content_type:
            self._send_json({"error": "Expected multipart/form-data"}, status=400)
            return

        content_length = int(self.headers.get("Content-Length", 0))
        body = self.rfile.read(content_length)

        fields, files = _parse_multipart(body, content_type)

        # Collect all uploaded files
        os.makedirs(TEMP_DIR, exist_ok=True)
        task_id = str(uuid.uuid4())[:8]
        _convert_progress[task_id] = {"status": "processing", "percent": 0}

        file_paths = []
        for fieldname, filename, data in files:
            if fieldname == "files" and filename:
                path = os.path.join(TEMP_DIR, f"{task_id}_{filename}")
                with open(path, "wb") as f:
                    f.write(data)
                file_paths.append(path)

        if len(file_paths) < 2:
            _convert_progress[task_id] = {"status": "error", "error": "Cần ít nhất 2 file PDF"}
            self._send_json({"task_id": task_id})
            return

        output_path = os.path.join(TEMP_DIR, f"{task_id}_merged.pdf")

        def _run():
            try:
                from PyPDF2 import PdfMerger
                _convert_progress[task_id] = {"status": "processing", "percent": 30}

                merger = PdfMerger()
                for p in file_paths:
                    merger.append(p)
                merger.write(output_path)
                merger.close()

                if os.path.exists(output_path):
                    file_id = str(uuid.uuid4())[:8]
                    _file_store[file_id] = output_path
                    _convert_progress[task_id] = {
                        "status": "done", "percent": 100,
                        "file_id": file_id, "file_name": "merged.pdf",
                        "file_size": os.path.getsize(output_path),
                    }
                else:
                    _convert_progress[task_id] = {"status": "error", "error": "Gộp PDF thất bại"}
            except ImportError:
                _convert_progress[task_id] = {"status": "error", "error": "Cần cài: pip install PyPDF2"}
            except Exception as e:
                _convert_progress[task_id] = {"status": "error", "error": str(e)}
            finally:
                for p in file_paths:
                    _cleanup_file(p)

        threading.Thread(target=_run, daemon=True).start()
        self._send_json({"task_id": task_id})

    def _handle_pdf_split(self):
        """Split a PDF into pages or ranges."""
        fields, file_data, file_name = self._read_multipart()
        if file_data is None:
            self._send_json({"error": "Không có file"}, status=400)
            return

        mode = fields.get("mode", "all")  # "all" or "ranges"
        ranges = fields.get("ranges", "")
        task_id = str(uuid.uuid4())[:8]
        _convert_progress[task_id] = {"status": "processing", "percent": 0}

        os.makedirs(TEMP_DIR, exist_ok=True)
        input_path = os.path.join(TEMP_DIR, f"{task_id}_{file_name}")
        with open(input_path, "wb") as f:
            f.write(file_data)

        def _run():
            try:
                from PyPDF2 import PdfReader, PdfWriter
                import zipfile

                _convert_progress[task_id] = {"status": "processing", "percent": 20}

                reader = PdfReader(input_path)
                total_pages = len(reader.pages)
                zip_path = os.path.join(TEMP_DIR, f"{task_id}_split.zip")
                base_name = os.path.splitext(file_name)[0]

                with zipfile.ZipFile(zip_path, "w") as zf:
                    if mode == "all":
                        for i, page in enumerate(reader.pages):
                            writer = PdfWriter()
                            writer.add_page(page)
                            page_path = os.path.join(TEMP_DIR, f"{task_id}_p{i+1}.pdf")
                            with open(page_path, "wb") as pf:
                                writer.write(pf)
                            zf.write(page_path, f"{base_name}_trang_{i+1}.pdf")
                            _cleanup_file(page_path)
                            _convert_progress[task_id] = {
                                "status": "processing",
                                "percent": 20 + int(70 * (i + 1) / total_pages),
                            }
                    else:
                        # Parse ranges like "1-3, 4-6, 7"
                        parts = [r.strip() for r in ranges.split(",") if r.strip()]
                        for idx, part in enumerate(parts):
                            writer = PdfWriter()
                            if "-" in part:
                                start, end = part.split("-", 1)
                                start, end = int(start) - 1, int(end)
                                for p in range(max(0, start), min(end, total_pages)):
                                    writer.add_page(reader.pages[p])
                                label = f"trang_{start+1}-{end}"
                            else:
                                p = int(part) - 1
                                if 0 <= p < total_pages:
                                    writer.add_page(reader.pages[p])
                                label = f"trang_{part}"

                            part_path = os.path.join(TEMP_DIR, f"{task_id}_part{idx}.pdf")
                            with open(part_path, "wb") as pf:
                                writer.write(pf)
                            zf.write(part_path, f"{base_name}_{label}.pdf")
                            _cleanup_file(part_path)

                if os.path.exists(zip_path):
                    file_id = str(uuid.uuid4())[:8]
                    _file_store[file_id] = zip_path
                    _convert_progress[task_id] = {
                        "status": "done", "percent": 100,
                        "file_id": file_id, "file_name": f"{base_name}_split.zip",
                        "file_size": os.path.getsize(zip_path),
                        "total_pages": total_pages,
                    }
                else:
                    _convert_progress[task_id] = {"status": "error", "error": "Tách PDF thất bại"}
            except ImportError:
                _convert_progress[task_id] = {"status": "error", "error": "Cần cài: pip install PyPDF2"}
            except Exception as e:
                _convert_progress[task_id] = {"status": "error", "error": str(e)}
            finally:
                _cleanup_file(input_path)

        threading.Thread(target=_run, daemon=True).start()
        self._send_json({"task_id": task_id})

    # ===== Speed Test Upload =====

    def _handle_speed_upload(self):
        """Receive upload data for speed test measurement."""
        content_length = int(self.headers.get("Content-Length", 0))
        # Just read and discard the data
        self.rfile.read(content_length)
        self._send_json({"success": True, "bytes": content_length})

    # ===== File Serve =====

    def _handle_file_serve(self, params):
        """Serve a file as browser download."""
        file_id = params.get("id", [""])[0]
        if not file_id or file_id not in _file_store:
            self._send_json({"error": "File không tồn tại hoặc đã hết hạn"}, status=404)
            return

        file_path = _file_store.pop(file_id)
        if not os.path.exists(file_path):
            self._send_json({"error": "File không tồn tại trên server"}, status=404)
            return

        file_name = os.path.basename(file_path)
        # Remove task_id prefix from filename
        if "_" in file_name and len(file_name.split("_")[0]) == 8:
            file_name = "_".join(file_name.split("_")[1:])

        file_size = os.path.getsize(file_path)
        content_type = mimetypes.guess_type(file_path)[0] or "application/octet-stream"

        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(file_size))
        # Use RFC 5987 encoding for non-ASCII filenames
        ascii_name = file_name.encode("ascii", errors="replace").decode("ascii")
        encoded_name = urllib.parse.quote(file_name)
        self.send_header("Content-Disposition",
                         f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{encoded_name}")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.end_headers()

        with open(file_path, "rb") as f:
            while chunk := f.read(65536):
                self.wfile.write(chunk)

        threading.Timer(5.0, lambda: _cleanup_file(file_path)).start()

    # ===== Common =====

    def _send_json(self, data, status=200):
        response = json.dumps(data, ensure_ascii=False, indent=2)
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()
        self.wfile.write(response.encode("utf-8"))

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def log_message(self, format, *args):
        print(f"[API] {args[0]}")


def _cleanup_file(path):
    try:
        if os.path.exists(path):
            os.remove(path)
    except OSError:
        pass


def main():
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    os.makedirs(TEMP_DIR, exist_ok=True)
    server = HTTPServer(("0.0.0.0", PORT), ToolsHandler)
    print(f"NQNghia Tools API running at http://localhost:{PORT}")
    print(f"Temp: {TEMP_DIR}")
    print()

    # Only open browser on local, not on production server
    if not os.environ.get("RENDER"):
        threading.Timer(1.0, lambda: webbrowser.open(f"http://localhost:{PORT}")).start()
        print(f"Opening browser at http://localhost:{PORT} ...")

    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nServer stopped.")
        server.server_close()


if __name__ == "__main__":
    main()
